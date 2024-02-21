from aiogram import types
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.builtin import CommandStart
from aiogram.types import CallbackQuery, InputFile, ReplyKeyboardMarkup, ReplyKeyboardRemove
from aiogram.dispatcher.filters import Regexp
import os

from data.config import ADMINS, text_about
from keyboards.default.check_menu import check_menu
from keyboards.inline.menu import menu, location, accept, play, menu_1
from loader import dp, bot

from docx import Document
from docx.shared import Inches

phone_regex = '^[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}$'


@dp.message_handler(CommandStart(), state='*')
async def bot_start(message: types.Message, state: FSMContext):
    await message.answer(f"Ҳурматли фойдаланувчи,сиз бу бот орқали бизнинг"
                         f" Електирик Мастер корхонамизга ариза юборишингиз мумкин!", reply_markup=menu)
    await state.finish()


@dp.callback_query_handler(text='about')
async def about_def(call: CallbackQuery, state: FSMContext):
    await call.message.edit_text(text=text_about, reply_markup=menu_1)


@dp.callback_query_handler(text='send')
async def an_def(call: CallbackQuery, state: FSMContext):
    await call.message.delete()
    await call.message.answer('Исм,Фамилянгизни киритинг!')
    await state.set_state('ism')


@dp.message_handler(state='ism')
async def ism_def(message: types.Message, state: FSMContext):
    ism = message.text
    await message.answer('Телефон ракамингизни киритинг.\n\nМисол учун:\n+998901234567\n 991234567')
    await state.update_data(
        {
            'ism': ism
        }
    )
    await state.set_state('phone')


@dp.message_handler(state='phone')
async def phone_def(message: types.Message, state: FSMContext):
    phone = message.text
    await state.update_data(
        {
            'phone': phone
        }
    )
    await message.answer(f'Пастдаги тугма орқали яшаш манзилингизни жўнатинг, (Location)', reply_markup=location)
    await state.set_state('location')


# @dp.message_handler(state='phone')
# async def space_def(message: types.Message):
#     await message.answer('🛑 Телефон ракамингизни киритинг!')


@dp.message_handler(state='location', content_types=['location'])
async def location_def(message: types.Location, state: FSMContext):
    full_location = message.values['location']
    lat = full_location['latitude']
    long = full_location['longitude']
    await state.update_data(
        {
            'longitude': long,
            'latitude': lat
        }
    )
    await bot.send_message(chat_id=message.chat.id,
                           text='Маълумотларингиз ва манзилингиз қабул қилинишига пастдаги тугма орқали рухсат беринг',
                           reply_markup=accept)
    await state.set_state('finish-state')


@dp.message_handler(state='finish-state')
async def finish_def(message: types.Message, state: FSMContext):
    data = await state.get_data()
    ism = data.get('ism')
    phone = data.get('phone')
    longitude = data.get('longitude')
    latitude = data.get('latitude')
    document = Document()

    document.add_heading('Malumotlar', 0)
    p = document.add_paragraph()
    p.add_run('Ma\'lumotlar maxfiy saqlanadi!').bold = True

    document.add_paragraph(f'Ism va familya: {ism} ', style='List Number')
    document.add_paragraph(f'Telefon raqami: {phone} ', style='List Number')
    document.add_page_break()
    document.save(f'{message.from_user.id}')

    doc = InputFile(f'{message.from_user.id}', filename=f'{message.from_user.full_name}.docx')
    await message.answer_document(document=doc, caption='Бу сизнинг аризангиз')
    await bot.send_message(chat_id=-1002039758691, text=f"Ismi:{ism}\n"
                                                        f"Telefon raqami:{phone}\n\n"
                                                        f"Pastdagi Location buyurtmachining manzili!👇👇👇👇")
    await bot.send_location(chat_id=-1002039758691, longitude=longitude, latitude=latitude)
    await message.answer('Маълумотларингиз қабул қилинди!\n\n'
                         '/start буйруги орқали ботни қайта ишга туширинг.', reply_markup=ReplyKeyboardRemove())
    await state.finish()
    os.remove(path=f'{message.from_user.id}')


def main():
    print('CI/CD')
